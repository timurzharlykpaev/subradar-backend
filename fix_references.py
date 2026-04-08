#!/usr/bin/env python3
"""
Скрипт для перенумерации ссылок в диссертации.
Двухпроходная замена: сначала все [N] -> [TEMPN], потом [TEMPN] -> [новый N].
Также переставляет источники в списке литературы.
"""

import re
import copy
from lxml import etree
from docx import Document

DOCX_PATH = "/Users/timurzharlykpaev/Desktop/repositories/subradar-backend/Жарлыкпаеваа Лаурраа проект магистратурааа.docx"
OUTPUT_PATH = "/Users/timurzharlykpaev/Desktop/repositories/subradar-backend/Жарлыкпаеваа Лаурраа проект магистратурааа_FIXED.docx"

# Маппинг: старый номер -> новый номер (по порядку первого появления в тексте)
OLD_TO_NEW = {
    1: 2,   2: 3,   3: 1,   4: 5,   5: 6,   6: 4,
    7: 7,   8: 8,   9: 9,  10: 10, 11: 11, 12: 12,
   13: 13, 14: 14, 15: 16, 16: 18, 17: 17, 18: 42,
   19: 19, 20: 20, 21: 15, 22: 21, 23: 22, 24: 23,
   25: 24, 26: 25, 27: 26, 28: 27, 29: 28, 30: 29,
   31: 30, 32: 32, 33: 40, 34: 31, 35: 44, 36: 36,
   37: 33, 38: 37, 39: 34, 40: 38, 41: 35, 42: 39,
   43: 41, 44: 43, 45: 45, 46: 46,
}

# Обратный маппинг: новый номер -> старый номер
NEW_TO_OLD = {v: k for k, v in OLD_TO_NEW.items()}

# Параграфы списка литературы: 553 (старый #1) до 598 (старый #46)
BIB_START = 553
BIB_END = 598  # inclusive


def replace_refs_in_text(text, mapping, phase):
    """
    phase 1: [N] -> [TEMP_N]
    phase 2: [TEMP_N] -> [new_N]
    """
    if phase == 1:
        def replacer(m):
            inner = m.group(1)
            parts = inner.split(',')
            new_parts = []
            for part in parts:
                part = part.strip()
                if part.isdigit():
                    new_parts.append(f"TEMP_{part}")
                else:
                    new_parts.append(part)
            return '[' + ', '.join(new_parts) + ']'
        return re.sub(r'\[(\d+(?:\s*,\s*\d+)*)\]', replacer, text)

    elif phase == 2:
        def replacer(m):
            inner = m.group(1)
            parts = inner.split(',')
            new_parts = []
            for part in parts:
                part = part.strip()
                if part.startswith('TEMP_'):
                    old_num = int(part.replace('TEMP_', ''))
                    new_num = mapping.get(old_num, old_num)
                    new_parts.append(str(new_num))
                else:
                    new_parts.append(part)
            # Sort numbers in ascending order within brackets
            try:
                new_parts_sorted = sorted(new_parts, key=lambda x: int(x))
                return '[' + ', '.join(new_parts_sorted) + ']'
            except ValueError:
                return '[' + ', '.join(new_parts) + ']'
        return re.sub(r'\[(TEMP_\d+(?:\s*,\s*TEMP_\d+)*)\]', replacer, text)


def process_paragraph_runs(paragraph, mapping, phase):
    """Process runs in a paragraph, handling cases where [N] spans multiple runs."""
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = replace_refs_in_text(full_text, mapping, phase)
    if new_text == full_text:
        return

    # Try per-run replacement first (preserves formatting)
    test_parts = []
    for run in paragraph.runs:
        test_parts.append(replace_refs_in_text(run.text, mapping, phase))

    if ''.join(test_parts) == new_text:
        for run in paragraph.runs:
            run.text = replace_refs_in_text(run.text, mapping, phase)
        return

    # References span multiple runs — merge into first run
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def process_table(table, mapping, phase):
    """Process all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph_runs(paragraph, mapping, phase)


def reorder_bibliography(doc):
    """
    Reorder bibliography by swapping paragraph XML elements.
    Entries use Word auto-numbering so position = number.
    """
    bib_paras = doc.paragraphs[BIB_START:BIB_END + 1]
    assert len(bib_paras) == 46, f"Expected 46 bib entries, got {len(bib_paras)}"

    # Verify content matches expected
    print(f"  Entry 1 (old): {bib_paras[0].text[:60]}...")
    print(f"  Entry 46 (old): {bib_paras[45].text[:60]}...")

    # Save deep copies of all paragraph XML elements
    # old_index 0 = old source #1, old_index 1 = old source #2, etc.
    xml_copies = []
    for para in bib_paras:
        xml_copies.append(copy.deepcopy(para._p))

    # Now replace: for new position i (0-based), we need old source NEW_TO_OLD[i+1]
    parent = bib_paras[0]._p.getparent()

    for new_pos_0based in range(46):
        new_num = new_pos_0based + 1  # 1-based new number
        old_num = NEW_TO_OLD[new_num]  # which old source goes here
        old_idx = old_num - 1  # 0-based index in xml_copies

        # Replace the XML element at this position
        target_para = bib_paras[new_pos_0based]
        old_p = target_para._p

        # Replace old_p with the copy
        new_p = copy.deepcopy(xml_copies[old_idx])
        parent.replace(old_p, new_p)

        # Update the reference in bib_paras so subsequent replacements work
        # (since parent.replace changes the tree)
        bib_paras[new_pos_0based]._p = new_p

    # Verify
    updated_paras = doc.paragraphs[BIB_START:BIB_END + 1]
    print(f"  Entry 1 (new): {updated_paras[0].text[:60]}...")
    print(f"  Entry 46 (new): {updated_paras[45].text[:60]}...")
    print("Bibliography reordered successfully")


def verify_references(doc):
    """Print all references found in text for verification."""
    ref_count = {}
    for para in doc.paragraphs:
        for m in re.finditer(r'\[(\d+(?:\s*,\s*\d+)*)\]', para.text):
            nums = [int(x.strip()) for x in m.group(1).split(',')]
            for n in nums:
                ref_count[n] = ref_count.get(n, 0) + 1

    print(f"\nReferences found in text: {sorted(ref_count.keys())}")
    missing = set(range(1, 47)) - set(ref_count.keys())
    if missing:
        print(f"WARNING: Missing references: {sorted(missing)}")
    extra = set(ref_count.keys()) - set(range(1, 47))
    if extra:
        print(f"WARNING: Extra references: {sorted(extra)}")

    # Check no TEMP_ remains
    for para in doc.paragraphs:
        if 'TEMP_' in para.text:
            print(f"WARNING: TEMP_ found in: {para.text[:80]}")


def main():
    print(f"Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    print("\n=== Phase 1: [N] -> [TEMP_N] ===")
    for para in doc.paragraphs:
        process_paragraph_runs(para, OLD_TO_NEW, phase=1)
    for table in doc.tables:
        process_table(table, OLD_TO_NEW, phase=1)

    print("=== Phase 2: [TEMP_N] -> [new_N] ===")
    for para in doc.paragraphs:
        process_paragraph_runs(para, OLD_TO_NEW, phase=2)
    for table in doc.tables:
        process_table(table, OLD_TO_NEW, phase=2)

    print("\n=== Phase 3: Reordering bibliography ===")
    reorder_bibliography(doc)

    print("\n=== Verification ===")
    verify_references(doc)

    print(f"\nSaving to: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("Done! Open the _FIXED file and check.")


if __name__ == '__main__':
    main()
